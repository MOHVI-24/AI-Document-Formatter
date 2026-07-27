"""Document parsing, optional AI enhancement, and export helpers."""
from __future__ import annotations

import html
import io
import json
import os
import re
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Literal

import httpx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

OutputFormat = Literal["pdf", "docx", "html"]
DocumentKind = Literal["auto", "report", "letter", "meeting-notes", "resume", "article"]


@dataclass
class StructuredDocument:
    title: str
    kind: str
    sections: list[tuple[str, list[str]]]
    summary: str
    enhanced: bool = False


def clean_text(text: str, max_chars: int) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").strip()
    if len(text) > max_chars:
        raise ValueError(f"Input is limited to {max_chars:,} characters.")
    if not text:
        raise ValueError("Add some text or upload a supported file before formatting.")
    return text


def extract_upload(filename: str, content: bytes) -> str:
    suffix = Path(filename or "upload.txt").suffix.lower()
    if suffix in {".txt", ".md"}:
        return content.decode("utf-8", errors="replace")
    if suffix == ".docx":
        try:
            return "\n".join(p.text for p in Document(io.BytesIO(content)).paragraphs if p.text.strip())
        except Exception as exc:  # malformed user file
            raise ValueError("That DOCX file could not be read.") from exc
    raise ValueError("Supported upload types are .txt, .md, and .docx.")


def detect_kind(text: str) -> str:
    lower = text.lower()
    if any(token in lower for token in ("dear ", "sincerely", "kind regards")):
        return "letter"
    if any(token in lower for token in ("attendees", "agenda", "action items", "minutes")):
        return "meeting-notes"
    if any(token in lower for token in ("experience", "education", "skills", "linkedin.com")):
        return "resume"
    if len(text.splitlines()) > 10 or any(token in lower for token in ("introduction", "conclusion", "methodology")):
        return "report"
    return "article"


def _title_from_text(text: str) -> str:
    first = next((line.strip("# -*\t ") for line in text.splitlines() if line.strip()), "Untitled document")
    if len(first) > 72 or first.endswith((".", "!", "?", ":")):
        words = re.findall(r"[A-Za-z0-9][A-Za-z0-9'/-]*", first)[:9]
        return " ".join(words).title() or "Untitled Document"
    return first.title()


def _split_sections(text: str) -> list[tuple[str, list[str]]]:
    sections: list[tuple[str, list[str]]] = []
    current_title, current = "Overview", []
    heading = re.compile(r"^(?:#{1,3}\s+|[A-Z][A-Za-z ]{2,}:$)")
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            if current and current[-1] != "":
                current.append("")
        elif heading.match(line):
            if current:
                sections.append((current_title, current))
            current_title, current = line.lstrip("#").rstrip(":").strip().title(), []
        else:
            current.append(line)
    if current:
        sections.append((current_title, current))
    return sections or [("Overview", [text])]


def offline_structure(text: str, selected_kind: DocumentKind) -> StructuredDocument:
    kind = detect_kind(text) if selected_kind == "auto" else selected_kind
    sections = _split_sections(text)
    body_words = " ".join(line for _, lines in sections for line in lines).split()
    summary = " ".join(body_words[:32]).rstrip(".,;:") + ("…" if len(body_words) > 32 else "")
    return StructuredDocument(_title_from_text(text), kind, sections, summary)


def _response_text(payload: dict) -> str:
    if payload.get("output_text"):
        return payload["output_text"]
    return "".join(content.get("text", "") for item in payload.get("output", []) for content in item.get("content", []) if content.get("type") == "output_text")


def enhance_with_openai(document: StructuredDocument, source_text: str) -> StructuredDocument:
    """Use the Responses API when configured; preserve local formatting on any failure."""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return document
    prompt = (
        "You are an exacting document editor. Improve grammar and clarity without adding facts. "
        "Return only JSON matching {title:string, kind:string, summary:string, sections:[{heading:string, paragraphs:[string]}]}. "
        "Keep paragraphs concise; retain lists using '- ' markers. Source text follows:\n\n" + source_text
    )
    payload = {"model": os.getenv("OPENAI_MODEL", "gpt-5.6-terra"), "reasoning": {"effort": "low"},
               "input": [{"role": "user", "content": [{"type": "input_text", "text": prompt}]}],
               "text": {"format": {"type": "json_object"}}}
    try:
        response = httpx.post("https://api.openai.com/v1/responses", headers={"Authorization": f"Bearer {api_key}"}, json=payload, timeout=45)
        response.raise_for_status()
        data = json.loads(_response_text(response.json()))
        sections = [(str(item["heading"]), [str(p) for p in item["paragraphs"]]) for item in data["sections"]]
        return StructuredDocument(str(data.get("title") or document.title), str(data.get("kind") or document.kind), sections,
                                  str(data.get("summary") or document.summary), True)
    except (httpx.HTTPError, KeyError, TypeError, ValueError, json.JSONDecodeError):
        return document


def build_document(text: str, kind: DocumentKind, use_ai: bool) -> StructuredDocument:
    document = offline_structure(text, kind)
    return enhance_with_openai(document, text) if use_ai else document


def _page_number(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#64748b"))
    canvas.drawRightString(A4[0] - 0.65 * inch, 0.42 * inch, str(doc.page))
    canvas.restoreState()


def export_pdf(document: StructuredDocument, accent: str) -> bytes:
    buffer, styles = io.BytesIO(), getSampleStyleSheet()
    title = ParagraphStyle("DocumentTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=25, leading=30, alignment=TA_CENTER, textColor=colors.HexColor(accent), spaceAfter=12)
    metadata = ParagraphStyle("Metadata", parent=styles["Normal"], fontSize=9, leading=13, alignment=TA_CENTER, textColor=colors.HexColor("#64748b"), spaceAfter=28)
    heading = ParagraphStyle("DocumentHeading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=15, leading=20, textColor=colors.HexColor(accent), spaceBefore=16, spaceAfter=7)
    body = ParagraphStyle("DocumentBody", parent=styles["BodyText"], fontSize=10.5, leading=16, spaceAfter=9)
    story = [Paragraph(html.escape(document.title), title), Paragraph(f"{html.escape(document.kind.replace('-', ' ').title())} · {datetime.now(UTC).strftime('%d %b %Y')}", metadata)]
    for index, (section_title, paragraphs) in enumerate(document.sections):
        if index and index % 8 == 0:
            story.append(PageBreak())
        story.append(Paragraph(html.escape(section_title), heading))
        for item in paragraphs:
            safe = html.escape(item)
            story.append(Paragraph(safe[2:] if item.startswith(("- ", "* ", "• ")) else safe or "&nbsp;", body, bulletText="•" if item.startswith(("- ", "* ", "• ")) else None))
        story.append(Spacer(1, 3))
    SimpleDocTemplate(buffer, pagesize=A4, rightMargin=.72*inch, leftMargin=.72*inch, topMargin=.78*inch, bottomMargin=.68*inch, title=document.title).build(story, onFirstPage=_page_number, onLaterPages=_page_number)
    return buffer.getvalue()


def export_docx(document: StructuredDocument, accent: str) -> bytes:
    doc = Document()
    doc.sections[0].top_margin = doc.sections[0].bottom_margin = Inches(.8)
    title = doc.add_heading(document.title, 0)
    title.runs[0].font.color.rgb = RGBColor.from_string(accent.lstrip("#"))
    doc.add_paragraph(f"{document.kind.replace('-', ' ').title()} · {datetime.now(UTC).strftime('%d %b %Y')}")
    for heading, paragraphs in document.sections:
        doc.add_heading(heading, 1)
        for item in paragraphs:
            doc.add_paragraph(item[2:] if item.startswith(("- ", "* ", "• ")) else item, style="List Bullet" if item.startswith(("- ", "* ", "• ")) else None)
    doc.styles["Normal"].font.size = Pt(11)
    buffer = io.BytesIO(); doc.save(buffer)
    return buffer.getvalue()


def export_html(document: StructuredDocument, accent: str) -> bytes:
    body = "".join(f"<section><h2>{html.escape(heading)}</h2>" + "".join(f"<p>{html.escape(item)}</p>" for item in paragraphs) + "</section>" for heading, paragraphs in document.sections)
    return f"<!doctype html><html lang=en><meta charset=utf-8><title>{html.escape(document.title)}</title><style>body{{font:17px/1.65 Georgia,serif;max-width:760px;margin:64px auto;color:#172033;padding:0 24px}}h1,h2{{font-family:Arial;color:{accent}}}h1{{font-size:2.7rem}}h2{{margin-top:2.4rem}}.meta{{color:#64748b}}</style><body><h1>{html.escape(document.title)}</h1><p class=meta>{html.escape(document.kind.title())}</p>{body}</body></html>".encode()


def export_document(document: StructuredDocument, output_format: OutputFormat, accent: str) -> tuple[bytes, str]:
    accent = accent if re.fullmatch(r"#[0-9a-fA-F]{6}", accent) else "#6d5dfc"
    if output_format == "pdf":
        return export_pdf(document, accent), "application/pdf"
    if output_format == "docx":
        return export_docx(document, accent), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return export_html(document, accent), "text/html; charset=utf-8"
